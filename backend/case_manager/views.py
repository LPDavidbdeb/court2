import html
import re
from datetime import datetime
from django.views.generic import ListView, DetailView, CreateView, View, FormView
from django.views.generic.edit import UpdateView
from django.urls import reverse_lazy, reverse
from django.shortcuts import redirect, get_object_or_404, render
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.conf import settings
from django.contrib import messages
from django.contrib.contenttypes.models import ContentType
from collections import defaultdict
import json
import google.generativeai as genai
import docx
import io
import os
import zipfile
from django.utils.html import strip_tags
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.views.decorators.http import require_POST

from .models import LegalCase, PerjuryContestation, AISuggestion, ExhibitRegistry, ProducedExhibit
from .forms import LegalCaseForm, PerjuryContestationForm, PerjuryContestationNarrativeForm, PerjuryContestationStatementsForm
from .services import refresh_case_exhibits, rebuild_produced_exhibits
from ai_services.utils import EvidenceFormatter
from ai_services.services import analyze_for_json_output, run_police_investigator_service, AI_PERSONAS
from document_manager.models import LibraryNode, DocumentSource, Statement
# NEW: Import rich document models for context lookup
from pdf_manager.models import PDFDocument
from email_manager.models import Email
from protagonist_manager.models import Protagonist

def _normalize_suggestion_json(data_dict):
    """
    Normalizes the AI suggestion JSON to a standard format.
    """
    normalized_data = {}
    
    # Find keys that seem to correspond to the four sections
    keys = list(data_dict.keys())
    
    # A common pattern is having 'suggestion_secX' and 'contenu_secX'
    # We prioritize 'contenu' if it exists
    for i in range(1, 5):
        title_key = f'suggestion_sec{i}'
        content_key = f'contenu_sec{i}'
        
        # Find the best key for the content
        content = data_dict.get(content_key, data_dict.get(title_key, ''))
        
        normalized_data[f'section_{i}'] = content

    return normalized_data

@require_POST
def update_contestation_title_ajax(request, pk):
    try:
        contestation = get_object_or_404(PerjuryContestation, pk=pk)
        data = json.loads(request.body)
        new_title = data.get('title', '').strip()

        if not new_title:
            return JsonResponse({'status': 'error', 'message': 'Title cannot be empty.'}, status=400)

        contestation.title = new_title
        contestation.save(update_fields=['title'])
        return JsonResponse({'status': 'success', 'new_title': new_title})

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON.'}, status=400)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def retry_parse_suggestion(request, suggestion_pk):
    suggestion = get_object_or_404(AISuggestion, pk=suggestion_pk)
    if not suggestion.raw_response:
        messages.error(request, "No raw response to parse.")
        return redirect('case_manager:contestation_detail', pk=suggestion.contestation.pk)

    try:
        json_match = re.search(r'\{.*\}', suggestion.raw_response, re.DOTALL)
        if json_match:
            cleaned_text = json_match.group(0)
            data_dict = json.loads(cleaned_text)
            
            suggestion.content = _normalize_suggestion_json(data_dict)
            suggestion.parsing_success = True
            suggestion.save()
            messages.success(request, "Successfully parsed and normalized the raw AI response.")
        else:
            messages.warning(request, "No JSON object could be found in the raw response.")

    except json.JSONDecodeError as e:
        messages.error(request, f"Failed to parse JSON: {e}")
    
    return redirect('case_manager:contestation_detail', pk=suggestion.contestation.pk)

def _get_allegation_context(case, targeted_statements):
    """
    Helper function to build the enriched text for allegations,
    grouping them by document and including the solemn declaration.
    """
    lies_text = "--- DÉCLARATIONS SOUS SERMENT (VERSION SUSPECTE) ---\n"
    statement_ids = [s.id for s in targeted_statements]

    # NEW: Create a lookup for rich document metadata
    rich_doc_metadata = {}
    # Get all PDF documents in the case
    pdf_docs = PDFDocument.objects.filter(quotes__trames_narratives__supported_contestations__case=case).select_related('author').distinct()
    for pdf in pdf_docs:
        rich_doc_metadata[pdf.title] = {'author': pdf.author, 'date': pdf.document_date}
    # Get all Emails in the case
    emails = Email.objects.filter(quotes__trames_narratives__supported_contestations__case=case).select_related('sender_protagonist').distinct()
    for email in emails:
        rich_doc_metadata[email.subject] = {'author': email.sender_protagonist, 'date': email.date_sent}

    stmt_content_type = ContentType.objects.get_for_model(Statement)
    nodes = LibraryNode.objects.filter(
        content_type=stmt_content_type,
        object_id__in=statement_ids,
        document__source_type=DocumentSource.REPRODUCED
    ).select_related('document', 'document__author')

    doc_to_stmts = defaultdict(list)
    for node in nodes:
        doc_to_stmts[node.document].append(node.content_object)

    for doc, stmts in doc_to_stmts.items():
        if doc.solemn_declaration:
            lies_text += f"CONTEXTE DU DOCUMENT : « {doc.title} »\n"
            lies_text += f"DÉCLARATION SOLENNELLE : « {doc.solemn_declaration} »\n\n"
        
        author_name = "Auteur Inconnu"
        author_role = ""
        doc_date = "Date Inconnue"

        if doc.source_type == DocumentSource.REPRODUCED:
            # Explicitly use Document model's fields for REPRODUCED documents
            if doc.author:
                author_name = doc.author.get_full_name()
                author_role = f" [{doc.author.role}]"
            if doc.document_original_date:
                doc_date = doc.document_original_date.strftime('%d %B %Y')
        else:
            # Existing logic for other document types (PDFDocument, Email) via rich_doc_metadata
            metadata = rich_doc_metadata.get(doc.title)
            if metadata:
                if metadata.get('author'):
                    author_name = metadata['author'].get_full_name()
                    author_role = f" [{metadata['author'].role}]"
                if metadata.get('date'):
                    doc_date = metadata['date'].strftime('%d %B %Y')
            elif doc.author: # Fallback to the generic document's author if not in rich_doc_metadata
                author_name = doc.author.get_full_name()
                author_role = f" [{doc.author.role}]"
            
            if not metadata and doc.document_original_date: # Fallback to generic date if not in rich_doc_metadata
                 doc_date = doc.document_original_date.strftime('%d %B %Y')

        for stmt in stmts:
            lies_text += f"[ {author_name}{author_role}, dans le document {doc.title} en date du {doc_date} ecrit : « {stmt.text} » ]\n\n"
    
    return lies_text

def serialize_evidence(evidence_pool):
    serialized_data = []
    if evidence_pool.get('events'):
        event_menu = []
        for event in evidence_pool['events']:
            event_menu.append({
                'title': f'Event: {event.date} - {event.explanation[:50]}...',
                'value': f'<blockquote><p>{event.explanation}</p><cite>Event on {event.date}</cite></blockquote>'
            })
        if event_menu:
            serialized_data.append({'title': 'Events', 'menu': event_menu})
    if evidence_pool.get('emails'):
        email_menu = []
        for quote in evidence_pool['emails']:
            email_menu.append({
                'title': f'Email Quote: {quote.quote_text[:50]}...',
                'value': f'<blockquote><p>{quote.quote_text}</p><cite>Email from {quote.email.sender} on {quote.email.date_sent}</cite></blockquote>'
            })
        if email_menu:
            serialized_data.append({'title': 'Email Quotes', 'menu': email_menu})
    return json.dumps(serialized_data)

def preview_ai_context(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    
    # 1. Refresh Registry & Data
    refresh_case_exhibits(contestation.case.pk)
    exhibit_registry = contestation.case.exhibits.all()
    exhibit_map = {
        (ex.content_type_id, ex.object_id): ex.get_label()
        for ex in exhibit_registry
    }

    evidence_data = EvidenceFormatter.collect_global_evidence(
        contestation.supporting_narratives.all()
    )

    # 2. Helper for natural sorting of exhibits (P-1, P-2, P-10...)
    def natural_keys(text):
        return [ int(c) if c.isdigit() else c for c in re.split(r'(\d+)', text) ]

    # 3. Build the XML Structure
    xml_output = []
    xml_output.append("<prompt_session>")

    # --- A. SYSTEM ROLE ---
    xml_output.append("""
    <system_role>
        You are a Senior Legal Strategist (Prosecutor) specializing in perjury demonstration.
        Your goal is to output a structured JSON strategy that proves the target made a false statement with INTENT (Mens Rea).
        You will act as a logic engine: analyze the <verified_facts> to disprove the <target_statement>.
    </system_role>
    """)

    # --- B. THE LIE (Target Statement) ---
    # We get the raw text and escape it to prevent XML breakage
    raw_allegation = _get_allegation_context(contestation.case, contestation.targeted_statements.all())
    xml_output.append(f"<target_statement>\n{html.escape(raw_allegation)}\n</target_statement>")

    # --- C. THE FACTS (Narrative Context) ---
    xml_output.append("<narrative_context>")
    for i, summary in enumerate(evidence_data['summaries'], 1):
        clean_summary = strip_tags(html.unescape(summary)).strip()
        xml_output.append(f"    <dimension id='{i}'>{html.escape(clean_summary)}</dimension>")
    xml_output.append("</narrative_context>")

    # --- D. THE TIMELINE (Chronology) ---
    xml_output.append("<verified_timeline>")
    for item in evidence_data['timeline']:
        obj_to_label = item.get('parent_doc', item['obj'])
        label = EvidenceFormatter.get_label(obj_to_label, exhibit_map)
        
        # We format the line using your utility, then escape it
        raw_line = EvidenceFormatter.format_timeline_item(item, exhibit_label=label)
        xml_output.append(f"    <event date='{item['date']}'>\n{html.escape(raw_line)}\n    </event>")
    xml_output.append("</verified_timeline>")

    # --- E. THE DOCUMENTS (Deep Dive) ---
    xml_output.append("<evidence_details>")
    sorted_docs = sorted(
        list(evidence_data['unique_documents']), 
        key=lambda d: natural_keys(EvidenceFormatter.get_label(d, exhibit_map) or "")
    )
    for doc in sorted_docs:
        label = EvidenceFormatter.get_label(doc, exhibit_map)
        label_str = f"P-{label}" if label else "NO-ID"
        
        # Get detailed content
        raw_content = EvidenceFormatter.format_document_reference(doc, exhibit_label=label)
        
        xml_output.append(f"    <document id='{label_str}'>")
        xml_output.append(f"<![CDATA[\n{raw_content}\n]]>") # CDATA handles newlines/special chars better for bulk text
        xml_output.append("    </document>")
    xml_output.append("</evidence_details>")

    # --- F. OUTPUT INSTRUCTIONS ---
    xml_output.append("""
    <task_instructions>
        <constraint>Do not output markdown formatting (like ```json). Return RAW JSON only.</constraint>
        <constraint>Tone: Clinical, factual, cold, surgical.</constraint>
        
        <output_schema>
        {
            "section_1": "The exact quote of the lie.",
            "section_2": "Summary of contrary facts (from <verified_timeline>).",
            "section_3": "Mens Rea Argument: Prove they KNEW it was false (e.g., 'Subject sent email P-12 contradicting this on date X').",
            "section_4": "Intent Argument: Why did they lie? (Judicial gain)."
        }
        </output_schema>
    </task_instructions>
    """)

    xml_output.append("</prompt_session>")

    return HttpResponse("\n".join(xml_output), content_type="text/plain; charset=utf-8")

def preview_police_prompt(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    narratives = contestation.supporting_narratives.all()
    
    xml_context = EvidenceFormatter.format_police_context_xml(narratives)
    
    persona = AI_PERSONAS['police_investigator']
    full_prompt = f"{persona['prompt']}\n\n{xml_context}"
    
    return HttpResponse(full_prompt, content_type="text/plain; charset=utf-8")

def generate_exhibit_production(request, pk):
    """
    Trigger to rebuild the ProducedExhibit table.
    """
    case = get_object_or_404(LegalCase, pk=pk)
    
    try:
        count = rebuild_produced_exhibits(case.pk)
        messages.success(request, f"Table des pièces générée avec succès ({count} entrées).")
    except Exception as e:
        messages.error(request, f"Erreur lors de la génération : {str(e)}")
        
    return redirect('case_manager:case_detail', pk=pk)

class LegalCaseListView(ListView):
    model = LegalCase
    template_name = 'case_manager/legalcase_list.html'
    context_object_name = 'cases'
    ordering = ['-created_at']

class LegalCaseDetailView(DetailView):
    model = LegalCase
    template_name = 'case_manager/legalcase_detail.html'
    context_object_name = 'case'
    def get(self, request, *args, **kwargs):
        refresh_case_exhibits(self.kwargs['pk'])
        return super().get(request, *args, **kwargs)
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['contestations'] = self.object.contestations.all()
        context['produced_exhibits'] = self.object.produced_exhibits.all()
        return context

class LegalCaseCreateView(CreateView):
    model = LegalCase
    form_class = LegalCaseForm
    template_name = 'case_manager/legalcase_form.html'
    def get_success_url(self):
        return reverse_lazy('case_manager:case_detail', kwargs={'pk': self.object.pk})

class LegalCaseExportView(View):
    def get(self, request, *args, **kwargs):
        case = get_object_or_404(LegalCase, pk=self.kwargs['pk'])
        
        # 1. RÉCUPÉRATION DE LA TABLE CALCULÉE (Source de vérité du site web)
        # On s'assure qu'elle est à jour avant l'export
        if not case.produced_exhibits.exists():
            rebuild_produced_exhibits(case.pk)
            
        produced_exhibits = ProducedExhibit.objects.filter(case=case).order_by('sort_order')
        
        document = docx.Document()

        def clean_text(text):
            if not text: return ""
            text = text.replace('</p>', '\n').replace('<br>', '\n').replace('<br/>', '\n')
            text = strip_tags(text)
            text = html.unescape(text)
            return text.strip()

        def add_hyperlink(paragraph, text, anchor):
            part = paragraph.part
            r_id = part.relate_to(anchor, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
            hyperlink.set(docx.oxml.shared.qn('w:anchor'), anchor, )
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            r = paragraph.add_run()
            r._r.append(hyperlink)
            r.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
            r.font.underline = True
            return hyperlink

        def add_markdown_content(doc, raw_text):
            # Version simplifiée sans renumbering_map car les labels sont fixes dans ProducedExhibit
            text = clean_text(raw_text)
            if not text: return

            text = re.sub(r'([\.\:\;])\s+([\*\-]\s)', r'\1\n\2', text)
            text = re.sub(r'([\.\:\;])\s+(\d+\.\s)', r'\1\n\2', text)
            lines = text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line: continue
                para_style = None
                if re.match(r'^[\*\-]\s+', line):
                    para_style = 'List Bullet'
                    line = re.sub(r'^[\*\-]\s+', '', line)
                elif re.match(r'^\d+\.\s+', line):
                    para_style = 'List Number'
                    line = re.sub(r'^[\*\-]\s+', '', line)

                p = doc.add_paragraph(style=para_style)
                p.add_run(line)

        # ------------------------------------------------------------------
        # DOCUMENT GENERATION
        # ------------------------------------------------------------------
        section = document.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        document.add_heading(f'Dénonciation: {case.title}', level=0)
        
        # --- SECTIONS ARGUMENTAIRES ---
        for contestation in case.contestations.all():
            document.add_heading(contestation.title, level=2)
            
            document.add_heading('1. Déclaration', level=3)
            add_markdown_content(document, contestation.final_sec1_declaration)
            
            document.add_heading('2. Preuve', level=3)
            add_markdown_content(document, contestation.final_sec2_proof)
            
            document.add_heading('3. Mens Rea', level=3)
            add_markdown_content(document, contestation.final_sec3_mens_rea)
            
            document.add_heading('4. Intention', level=3)
            add_markdown_content(document, contestation.final_sec4_intent)
            
            document.add_page_break()

        # ==================================================================
        # TABLE DES PIÈCES (Format Site Web)
        # ==================================================================
        document.add_heading('Index des Pièces (Production)', level=1)
        
        # Création du tableau à 5 colonnes
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False 
        
        # En-têtes
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cote'
        hdr_cells[0].width = Inches(0.8)
        
        hdr_cells[1].text = 'Date'
        hdr_cells[1].width = Inches(1.0)
        
        hdr_cells[2].text = 'Type'
        hdr_cells[2].width = Inches(1.0)
        
        hdr_cells[3].text = 'Description'
        hdr_cells[3].width = Inches(3.0)
        
        hdr_cells[4].text = 'Parties'
        hdr_cells[4].width = Inches(1.5)

        # Remplissage avec ProducedExhibit
        for item in produced_exhibits:
            row_cells = table.add_row().cells
            
            # Cell 0: Cote avec lien interne vers l'annexe
            bookmark_name = f"exhibit_{item.sort_order}" # ex: exhibit_1, exhibit_2
            add_hyperlink(row_cells[0].paragraphs[0], item.label, bookmark_name)
            
            # Cell 1: Date
            row_cells[1].text = item.date_display or ""
            
            # Cell 2: Type
            row_cells[2].text = item.exhibit_type or ""
            
            # Cell 3: Description (Nettoyage léger)
            desc_clean = clean_text(item.description)
            # Pour les citations, on peut mettre en italique
            if "«" in desc_clean:
                row_cells[3].paragraphs[0].add_run(desc_clean).italic = True
            else:
                row_cells[3].text = desc_clean
            
            # Cell 4: Parties
            row_cells[4].text = item.parties or ""

        # ==================================================================
        # ANNEXES (Basé sur ProducedExhibit)
        # ==================================================================
        document.add_page_break()
        document.add_heading('ANNEXES - CONTENU DÉTAILLÉ', level=0)

        for item in produced_exhibits:
            obj = item.content_object # L'objet réel (Email, PDF, etc.)
            if not obj: continue # Sécurité si l'objet a été supprimé

            label = item.label
            bookmark_name = f"exhibit_{item.sort_order}"
            
            # Heading avec Bookmark
            heading_paragraph = document.add_heading(f'Pièce {label}', level=1)
            bookmark_start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
            bookmark_start.set(docx.oxml.shared.qn('w:id'), str(item.sort_order))
            bookmark_start.set(docx.oxml.shared.qn('w:name'), bookmark_name)
            heading_paragraph._p.insert(0, bookmark_start)
            bookmark_end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
            bookmark_end.set(docx.oxml.shared.qn('w:id'), str(item.sort_order))
            heading_paragraph._p.append(bookmark_end)

            # --- Affichage conditionnel selon le type d'objet ---
            # On utilise item.content_type.model pour savoir comment l'afficher
            model_name = item.content_type.model

            if model_name == 'email' or model_name == 'quote': 
                # Note: 'quote' pointe souvent vers un Email ou un PDF, il faut gérer le parent
                actual_obj = obj
                if model_name == 'quote':
                    if hasattr(obj, 'email'): actual_obj = obj.email
                    elif hasattr(obj, 'pdf_document'): actual_obj = obj.pdf_document
                
                # Affiche le contexte de base
                p = document.add_paragraph()
                p.add_run(f"Description : {item.description}\n").bold = True
                p.add_run(f"Parties : {item.parties}").italic = True
                
                document.add_paragraph('--- Contenu ---').italic = True
                
                if hasattr(actual_obj, 'body_plain_text'):
                    raw_body = actual_obj.body_plain_text or "[Vide]"
                    # Restore email history stripping
                    body_lines = raw_body.splitlines()
                    cleaned_lines = [line for line in body_lines if not line.strip().startswith('>')]
                    cleaned_body = "\n".join(cleaned_lines)
                    body_text = clean_text(cleaned_body)
                    document.add_paragraph(body_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            elif model_name == 'event':
                document.add_paragraph(f"Date : {obj.date}")
                p = document.add_paragraph()
                p.add_run("Description : ").bold = True
                add_markdown_content(document, obj.explanation)

                # Restore photo display for events
                photos = obj.linked_photos.all()
                if photos.exists():
                    document.add_paragraph("Preuve visuelle :").italic = True
                    photo_table = document.add_table(rows=0, cols=2)
                    row_cells = None
                    for index, photo in enumerate(photos):
                        if index % 2 == 0:
                            row_cells = photo_table.add_row().cells
                        cell = row_cells[index % 2]
                        if photo.file:
                            try:
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(photo.file.open(), width=Inches(2.8))
                                caption = cell.add_paragraph(photo.file_name or "Image")
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                cell.add_paragraph(f"[Erreur: {e}]")

            elif model_name == 'photodocument':
                document.add_paragraph(f"Titre : {obj.title}")
                if obj.description:
                    add_markdown_content(document, obj.description)
                
                # Restore photo display for photodocuments
                photos = obj.photos.all()
                if photos.exists():
                    photo_table = document.add_table(rows=0, cols=2)
                    row_cells = None
                    for index, photo in enumerate(photos):
                        if index % 2 == 0:
                            row_cells = photo_table.add_row().cells
                        cell = row_cells[index % 2]
                        if photo.file:
                            try:
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(photo.file.open(), width=Inches(2.8))
                                caption = cell.add_paragraph(f"Page {index + 1}")
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                cell.add_paragraph(f"[Erreur: {e}]")

            elif model_name == 'pdfdocument' or model_name == 'document':
                document.add_paragraph(f"Document : {item.description}")
                document.add_paragraph(f"Auteur : {item.parties}")
                document.add_paragraph("[Voir fichier PDF joint au dossier]").italic = True

            elif model_name == 'statement':
                 document.add_paragraph(f"Déclaration : {obj.text}")

            document.add_page_break()

        # Save
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.getvalue(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="case_{case.pk}_export_v2.docx"'
        return response

class LegalCaseLLMExportView(View):
    def get(self, request, *args, **kwargs):
        case = get_object_or_404(LegalCase, pk=self.kwargs['pk'])
        
        # Ensure the table is up to date
        if not case.produced_exhibits.exists():
            rebuild_produced_exhibits(case.pk)
            
        produced_exhibits = ProducedExhibit.objects.filter(case=case).order_by('sort_order')
        
        # Build Markdown Content
        md_lines = []
        md_lines.append(f"# EVIDENCE TABLE: {case.title}")
        md_lines.append(f"Generated: {timezone.now().strftime('%Y-%m-%d')}")
        md_lines.append("Format: Markdown Key-Value (Optimized for LLM Context)")
        md_lines.append("---")
        
        for item in produced_exhibits:
            # Header with ID
            md_lines.append(f"## {item.label}")
            
            # Key-Value Metadata
            md_lines.append(f"- **Date**: {item.date_display or 'Unknown'}")
            md_lines.append(f"- **Type**: {item.exhibit_type or 'Unknown'}")
            md_lines.append(f"- **Parties**: {item.parties or 'Unknown'}")
            
            # Description / Content
            # Clean up description to be single-line or blockquote if needed
            desc = item.description.replace('\n', ' ').strip()
            md_lines.append(f"- **Description**: {desc}")
            
            # Add content from the object if available (e.g. email body)
            obj = item.content_object
            if obj:
                model_name = item.content_type.model
                content_text = ""
                
                if model_name in ['email', 'quote']:
                    actual_obj = obj
                    if model_name == 'quote':
                        if hasattr(obj, 'email'): actual_obj = obj.email
                        elif hasattr(obj, 'pdf_document'): actual_obj = obj.pdf_document
                    
                    if hasattr(actual_obj, 'body_plain_text') and actual_obj.body_plain_text:
                        # Simple cleanup
                        body = actual_obj.body_plain_text
                        # Remove reply chains roughly
                        lines = [l for l in body.splitlines() if not l.strip().startswith('>')]
                        content_text = "\n".join(lines).strip()
                        
                elif model_name == 'statement':
                    content_text = obj.text
                    
                elif model_name == 'event':
                    content_text = obj.explanation
                
                if content_text:
                    # Truncate if too long to save tokens, or keep full? 
                    # For LLM context, full is usually better unless massive.
                    # Let's limit to ~2000 chars for safety in this specific view
                    if len(content_text) > 2000:
                        content_text = content_text[:2000] + "... [TRUNCATED]"
                    
                    md_lines.append(f"- **Content**: \"{content_text}\"")
            
            md_lines.append("") # Empty line between items
            
        # Create Response
        response_text = "\n".join(md_lines)
        response = HttpResponse(response_text, content_type='text/markdown; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="Case_{case.pk}_LLM_Context.md"'
        return response

def download_exhibits_zip(request, pk):
    """
    Generates a ZIP file containing all produced exhibits for a case.
    Files are renamed with their exhibit label (e.g., 'P-1_contract.pdf').
    """
    case = get_object_or_404(LegalCase, pk=pk)
    
    # Ensure the table is up to date
    if not case.produced_exhibits.exists():
        from .services import rebuild_produced_exhibits
        rebuild_produced_exhibits(case.pk)

    exhibits = ProducedExhibit.objects.filter(case=case).order_by('sort_order')

    # Create an in-memory byte buffer to hold the zip file
    buffer = io.BytesIO()

    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for exhibit in exhibits:
            obj = exhibit.content_object
            if not obj:
                continue

            model_name = exhibit.content_type.model
            
            # --- 1. HANDLE PDF DOCUMENTS ---
            if model_name == 'pdfdocument':
                if obj.file and hasattr(obj.file, 'path') and os.path.exists(obj.file.path):
                    try:
                        with obj.file.open('rb') as f:
                            _, ext = os.path.splitext(obj.file.name)
                            safe_title = "".join([c for c in obj.title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
                            file_name = f"{exhibit.label}_{safe_title.replace(' ', '_')}{ext}"
                            zip_file.writestr(file_name, f.read())
                    except Exception as e:
                        print(f"Could not add PDF {exhibit.label} to zip: {e}")
                continue

            # --- 2. HANDLE EMAILS ---
            elif model_name == 'email':
                file_added = False
                # Priority 1: The FileField, if the file exists
                if obj.eml_file and hasattr(obj.eml_file, 'path') and os.path.exists(obj.eml_file.path):
                    try:
                        with obj.eml_file.open('rb') as f:
                            file_name = f"{exhibit.label}_{obj.subject[:50].replace(' ', '_')}.eml"
                            zip_file.writestr(file_name, f.read())
                            file_added = True
                    except Exception as e:
                        print(f"Error reading email from FileField for exhibit {exhibit.label}: {e}")
                
                # Priority 2: The raw file path, if the first method failed and the path exists
                if not file_added and obj.eml_file_path and os.path.exists(obj.eml_file_path):
                    try:
                        file_name = f"{exhibit.label}_{obj.subject[:50].replace(' ', '_')}.eml"
                        zip_file.write(obj.eml_file_path, arcname=file_name)
                    except Exception as e:
                        print(f"Error reading email from eml_file_path for exhibit {exhibit.label}: {e}")
                continue

            # --- 3. HANDLE PHOTO DOCUMENTS (Container of photos) ---
            elif model_name == 'photodocument':
                for i, photo in enumerate(obj.photos.all(), 1):
                    if photo.file and hasattr(photo.file, 'path') and os.path.exists(photo.file.path):
                        try:
                            with photo.file.open('rb') as f:
                                _, ext = os.path.splitext(photo.file.name)
                                photo_name = f"{exhibit.label}_{i:02d}_{obj.title[:30].replace(' ', '_')}{ext}"
                                zip_file.writestr(photo_name, f.read())
                        except Exception as e:
                            print(f"Error reading photo {photo.id} for exhibit {exhibit.label}: {e}")
                continue

            # --- 4. HANDLE EVENT (with linked photos) ---
            elif model_name == 'event':
                for i, photo in enumerate(obj.linked_photos.all(), 1):
                    if photo.file and hasattr(photo.file, 'path') and os.path.exists(photo.file.path):
                        try:
                            with photo.file.open('rb') as f:
                                _, ext = os.path.splitext(photo.file.name)
                                # Use the photo's own file_name if available, fallback to event explanation
                                safe_name = photo.file_name or obj.explanation[:30]
                                photo_name = f"{exhibit.label}_{i:02d}_{safe_name.replace(' ', '_')}{ext}"
                                zip_file.writestr(photo_name, f.read())
                        except Exception as e:
                            print(f"Error reading photo {photo.id} for event {exhibit.label}: {e}")
                continue

            # =========================================================
            # 🆕 ADD THIS BLOCK FOR LINKED DOCUMENTS
            # =========================================================
            elif model_name == 'document':
                # Check if the document has a file_source attached
                if obj.file_source and hasattr(obj.file_source, 'path') and os.path.exists(obj.file_source.path):
                    try:
                        with obj.file_source.open('rb') as f:
                            # 1. Determine Extension (Default to .pdf if missing)
                            _, ext = os.path.splitext(obj.file_source.name)
                            if not ext: 
                                ext = ".pdf"
                            
                            # 2. Create a clean filename for the Zip
                            # Format: "P-12_Title_of_Document.pdf"
                            safe_title = "".join([c for c in obj.title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
                            safe_title = safe_title.replace(' ', '_')
                            file_name = f"{exhibit.label}_{safe_title}{ext}"
                            
                            # 3. Write to Zip
                            zip_file.writestr(file_name, f.read())
                            
                    except Exception as e:
                        print(f"Could not add Document {exhibit.label} to zip: {e}")
                continue
            # =========================================================

    # Finalize the zip
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="Exhibits_{case.title}_Export.zip"'
    
    return response

class PoliceComplaintExportView(View):
    def get(self, request, *args, **kwargs):
        case = get_object_or_404(LegalCase, pk=self.kwargs['pk'])
        document = docx.Document()
        
        style = document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        document.add_heading(f"DOSSIER DE PLAINTE : {case.title}", level=0)
        document.add_paragraph(f"Date du rapport : {timezone.now().strftime('%Y-%m-%d')}")
        document.add_paragraph("À l'attention des enquêteurs.")
        document.add_page_break()

        contestations = case.contestations.exclude(police_report_data={})

        if not contestations.exists():
            document.add_paragraph("Aucune plainte policière n'a été générée pour ce dossier.")

        for contestation in contestations:
            data = contestation.police_report_data
            
            titre = data.get('titre_document', f"PLAINTE - {contestation.title}")
            document.add_heading(titre, level=1)
            
            sections = data.get('sections', [])
            for section in sections:
                if 'titre' in section:
                    document.add_heading(section['titre'], level=2)
                
                if 'contenu' in section:
                    content = section['contenu']
                    if isinstance(content, list):
                        for item in content:
                            p = document.add_paragraph(str(item))
                            p.style = 'List Bullet'
                    else:
                        document.add_paragraph(str(content))
            
            document.add_page_break()

        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="PLAINTE_POLICE_{case.pk}.docx"'
        return response

class PerjuryContestationCreateView(CreateView):
    model = PerjuryContestation
    form_class = PerjuryContestationForm
    template_name = 'case_manager/perjurycontestation_form.html'
    def form_valid(self, form):
        form.instance.case = get_object_or_404(LegalCase, pk=self.kwargs['case_pk'])
        return super().form_valid(form)
    def get_success_url(self):
        return reverse_lazy('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

class ManageContestationNarrativesView(UpdateView):
    model = PerjuryContestation
    form_class = PerjuryContestationNarrativeForm
    template_name = 'case_manager/manage_narratives.html'
    context_object_name = 'contestation'

    def get_success_url(self):
        messages.success(self.request, "Supporting narratives updated successfully.")
        return reverse('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

class ManageContestationStatementsView(UpdateView):
    model = PerjuryContestation
    form_class = PerjuryContestationStatementsForm
    template_name = 'case_manager/manage_statements.html'
    context_object_name = 'contestation'

    def get_success_url(self):
        messages.success(self.request, "Targeted statements updated successfully.")
        return reverse('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

class PerjuryContestationDetailView(UpdateView):
    model = PerjuryContestation
    template_name = 'case_manager/perjurycontestation_detail.html'
    context_object_name = 'contestation'
    fields = ['final_sec1_declaration', 'final_sec2_proof', 'final_sec3_mens_rea', 'final_sec4_intent']
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        evidence_pool = {'events': [], 'emails': [], 'pdfs': []}
        for narrative in self.object.supporting_narratives.all():
            evidence_pool['events'].extend(narrative.evenements.all())
            evidence_pool['emails'].extend(narrative.citations_courriel.all())
            evidence_pool['pdfs'].extend(narrative.citations_pdf.all())
        context['evidence_json'] = serialize_evidence(evidence_pool)
        context['ai_drafts'] = self.object.ai_suggestions.order_by('-created_at')
        return context

    def get_success_url(self):
        return reverse('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

def generate_ai_suggestion(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    
    narratives_context = []
    
    for narrative in contestation.supporting_narratives.all():
        analysis = narrative.get_structured_analysis()
        
        narrative_block = f"--- TRAME FACTUELLE : {narrative.titre} ---\n"
        
        if 'constats_objectifs' in analysis:
            for constat in analysis['constats_objectifs']:
                narrative_block += f"FAIT ÉTABLI : {constat.get('fait_identifie', 'N/A')}\n"
                narrative_block += f"DÉTAIL : {constat.get('description_factuelle', '')}\n"
                narrative_block += f"IMPACT : {constat.get('contradiction_directe', '')}\n\n"
        else:
            narrative_block += f"RÉSUMÉ MANUEL : {narrative.resume}\n"
            
        narratives_context.append(narrative_block)

    full_evidence_text = "\n".join(narratives_context)

    prompt_sequence = [
        """
        RÔLE : Stratège Juridique Senior (Procureur).
        MISSION : Rédiger un argumentaire de parjure dévastateur basé sur des FAITS VÉRIFIÉS.
        
        TU NE DOIS PAS : Chercher des preuves (c'est déjà fait).
        TU DOIS : Prouver l'INTENTION de mentir (Mens Rea) en connectant les faits.
        """,
        
        f"=== CIBLE (DÉCLARATION SOUS SERMENT) ===\n{_get_allegation_context(contestation.case, contestation.targeted_statements.all())}",
        
        f"=== AUDIT DES FAITS (PREUVE IRRÉFUTABLE) ===\n{full_evidence_text}",
        
        """
        === DIRECTIVES DE RÉDACTION ===
        Rédige le rapport au format JSON strict.
        
        Section 3 (Mens Rea) est la plus importante : Explique comment la multiplicité des faits (les dates, les photos, les emails) prouve qu'il est IMPOSSIBLE que le sujet ait fait une simple "erreur". C'est un mensonge calculé.
        
        Structure JSON attendue :
        {
            "section_1": "Citation exacte et contexte...",
            "section_2": "Synthèse des faits contraires (utilise les faits de l'audit)...",
            "section_3": "Argumentaire sur la Connaissance (Mens Rea)...",
            "section_4": "Argumentaire sur l'Intention (Gain judiciaire)..."
        }
        """
    ]
    
    try:
        raw_text = analyze_for_json_output(prompt_sequence)
        
        suggestion = AISuggestion.objects.create(
            contestation=contestation,
            raw_response=raw_text,
            content={},
            parsing_success=False
        )

        try:
            data_dict = json.loads(raw_text)
            suggestion.content = _normalize_suggestion_json(data_dict)
            suggestion.parsing_success = True
            suggestion.save()
            messages.success(request, "Stratégie de parjure générée avec succès sur la base de l'audit.")
        except json.JSONDecodeError:
            messages.warning(request, "Réponse générée mais format JSON invalide.")

    except Exception as e:
        messages.error(request, f"Erreur API : {e}")

    return redirect('case_manager:contestation_detail', pk=contestation.pk)

def generate_police_report(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    
    narratives = contestation.supporting_narratives.prefetch_related(
        'evenements', 'citations_courriel', 'citations_pdf', 'photo_documents', 'targeted_statements'
    ).all()
    
    try:
        raw_json = run_police_investigator_service(narratives)
        data = json.loads(raw_json)
        
        contestation.police_report_data = data
        contestation.police_report_date = timezone.now()
        contestation.save()
        
        messages.success(request, "Rapport de police généré avec succès.")
    except Exception as e:
        messages.error(request, f"Erreur lors de la génération : {e}")
        
    return redirect('case_manager:contestation_detail', pk=contestation.pk)

def case_protagonists_list(request, pk):
    case = get_object_or_404(LegalCase, pk=pk)
    
    # Ensure exhibits are up to date
    if not case.produced_exhibits.exists():
        rebuild_produced_exhibits(case.pk)
    
    # Collect all distinct protagonists from the exhibits
    protagonists = set()
    
    # Iterate through all produced exhibits
    for exhibit in case.produced_exhibits.all():
        obj = exhibit.content_object
        if not obj:
            continue
            
        model_name = exhibit.content_type.model
        
        if model_name == 'email':
            if obj.sender_protagonist:
                protagonists.add(obj.sender_protagonist)
            for recipient in obj.recipient_protagonists.all():
                protagonists.add(recipient)
                
        elif model_name == 'pdfdocument':
            if obj.author:
                protagonists.add(obj.author)
                
        elif model_name == 'document':
            if obj.author:
                protagonists.add(obj.author)
                
        elif model_name == 'photodocument':
            if obj.author:
                protagonists.add(obj.author)
                
        elif model_name == 'chatsequence':
            # For chat sequences, we need to check the messages
            for msg in obj.messages.all():
                if msg.sender and msg.sender.protagonist:
                    protagonists.add(msg.sender.protagonist)
    
    # Convert set to list and sort by name
    protagonists_list = sorted(list(protagonists), key=lambda p: p.get_full_name())
    
    context = {
        'case': case,
        'protagonists': protagonists_list
    }
    
    return render(request, 'case_manager/case_protagonists.html', context)
