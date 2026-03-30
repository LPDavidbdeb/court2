import docx
import io
import re
import html
from django.utils.html import strip_tags
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import LegalCase, ProducedExhibit, PerjuryContestation
from .services import rebuild_produced_exhibits

def clean_text(text: str) -> str:
    if not text: return ""
    text = text.replace('</p>', '\n').replace('<br>', '\n').replace('<br/>', '\n')
    text = strip_tags(text)
    text = html.unescape(text)
    return text.strip()

def add_hyperlink(paragraph, text, anchor):
    part = paragraph.part
    r_id = part.relate_to(anchor, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), anchor, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
    r.font.underline = True
    return hyperlink

def add_markdown_content(doc, raw_text):
    text = clean_text(raw_text)
    if not text: return

    text = re.sub(r'([\.\:\;])\s+([\*\-]\s)', r'\1\n\2', text)
    text = re.sub(r'([\.\:\;])\s+(\d+\.\s)', r'\1\n\2', text)
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line: continue
        para_style = None
        if re.match(r'^[\*\-]\s+', line):
            para_style = 'List Bullet'
            line = re.sub(r'^[\*\-]\s+', '', line)
        elif re.match(r'^\d+\.\s+', line):
            para_style = 'List Number'
            line = re.sub(r'^[\*\-]\s+', '', line)

        p = doc.add_paragraph(style=para_style)
        p.add_run(line)

def generate_case_docx(case_id: int) -> io.BytesIO:
    case = LegalCase.objects.get(pk=case_id)
    
    if not case.produced_exhibits.exists():
        rebuild_produced_exhibits(case.pk)
            
    produced_exhibits = ProducedExhibit.objects.filter(case=case).order_by('sort_order')
    
    document = docx.Document()
    section = document.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    document.add_heading(f'Dénonciation: {case.title}', level=0)
    
    for contestation in case.contestations.all():
        document.add_heading(contestation.title, level=2)
        
        document.add_heading('1. Déclaration', level=3)
        add_markdown_content(document, contestation.final_sec1_declaration)
        
        document.add_heading('2. Preuve', level=3)
        add_markdown_content(document, contestation.final_sec2_proof)
        
        document.add_heading('3. Mens Rea', level=3)
        add_markdown_content(document, contestation.final_sec3_mens_rea)
        
        document.add_heading('4. Intention', level=3)
        add_markdown_content(document, contestation.final_sec4_intent)
        
        document.add_page_break()

    # Index Table
    document.add_heading('Index des Pièces (Production)', level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cote'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Type'
    hdr_cells[3].text = 'Description'
    hdr_cells[4].text = 'Parties'

    for item in produced_exhibits:
        row_cells = table.add_row().cells
        bookmark_name = f"exhibit_{item.sort_order}"
        add_hyperlink(row_cells[0].paragraphs[0], item.label, bookmark_name)
        row_cells[1].text = item.date_display or ""
        row_cells[2].text = item.exhibit_type or ""
        desc_clean = clean_text(item.description)
        if "«" in desc_clean:
            row_cells[3].paragraphs[0].add_run(desc_clean).italic = True
        else:
            row_cells[3].text = desc_clean
        row_cells[4].text = item.parties or ""

    # Annexes
    document.add_page_break()
    document.add_heading('ANNEXES - CONTENU DÉTAILLÉ', level=0)

    for item in produced_exhibits:
        obj = item.content_object
        if not obj: continue

        label = item.label
        bookmark_name = f"exhibit_{item.sort_order}"
        
        heading_paragraph = document.add_heading(f'Pièce {label}', level=1)
        # Bookmark logic here (omitted for brevity in this example but should be included)
        
        model_name = item.content_type.model
        # ... Add detailed content logic from views.py ...
        
        document.add_page_break()

    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f
